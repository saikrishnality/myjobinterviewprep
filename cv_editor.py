#!/usr/bin/env python3
"""
cv_editor.py — Edits a Word CV template using python-docx
Usage: python3 cv_editor.py <input.docx> <output.docx> <cvdata.json> [tracked]
"""

import sys
import json
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime, timezone

DATE = datetime.now(timezone.utc).strftime('%Y-%m-%dT%H:%M:%SZ')
AUTHOR = 'CV Match'
_id_counter = [300]

def next_id():
    _id_counter[0] += 1
    return str(_id_counter[0])

def get_text(para):
    return para.text.strip()

# ── CLEAN REPLACEMENT ─────────────────────────────────────────
def replace_clean(para, new_text):
    runs = para.runs
    if runs:
        runs[0].text = new_text
        for run in runs[1:]:
            run.text = ''
    else:
        para.add_run(new_text)

def clear_para(para):
    for run in para.runs:
        run.text = ''

# ── TRACKED REPLACEMENT ───────────────────────────────────────
def replace_tracked(para, old_text, new_text):
    if not old_text.strip() or old_text.strip() == new_text.strip():
        replace_clean(para, new_text)
        return
    p_elem = para._p
    for child in list(p_elem):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag in ('r', 'ins', 'del', 'hyperlink'):
            p_elem.remove(child)
    del_elem = OxmlElement('w:del')
    del_elem.set(qn('w:id'), next_id())
    del_elem.set(qn('w:author'), AUTHOR)
    del_elem.set(qn('w:date'), DATE)
    r_del = OxmlElement('w:r')
    dt = OxmlElement('w:delText')
    dt.set(qn('xml:space'), 'preserve')
    dt.text = old_text
    r_del.append(dt)
    del_elem.append(r_del)
    p_elem.append(del_elem)
    ins_elem = OxmlElement('w:ins')
    ins_elem.set(qn('w:id'), next_id())
    ins_elem.set(qn('w:author'), AUTHOR)
    ins_elem.set(qn('w:date'), DATE)
    r_ins = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = new_text
    r_ins.append(t)
    ins_elem.append(r_ins)
    p_elem.append(ins_elem)

def replace(para, old_text, new_text, tracked):
    if tracked:
        replace_tracked(para, old_text, new_text)
    else:
        replace_clean(para, new_text)

# ── MAIN EDITOR ───────────────────────────────────────────────
def edit_cv(input_path, output_path, cv_data, tracked=False):
    doc = Document(input_path)
    paras = doc.paragraphs

    # Find section boundaries
    summary_start = expertise_idx = experience_idx = None
    for i, para in enumerate(paras):
        text = get_text(para).upper()
        if 'PROFESSIONAL SUMMARY' in text and summary_start is None:
            summary_start = i
        elif 'EXPERTISE' in text and summary_start is not None and expertise_idx is None:
            expertise_idx = i
        elif 'PROFESSIONAL EXPERIENCE' in text and expertise_idx is not None and experience_idx is None:
            experience_idx = i

    print(f"Sections: summary={summary_start} expertise={expertise_idx} experience={experience_idx}", file=sys.stderr)

    # Replace Professional Summary
    if summary_start is not None and expertise_idx is not None:
        new_parts = [p.strip() for p in cv_data.get('summary','').split('\n\n') if p.strip()]
        if not new_parts:
            new_parts = [cv_data.get('summary','')]
        sum_paras = [p for p in paras[summary_start+1:expertise_idx] if get_text(p)]
        for i, para in enumerate(sum_paras):
            old_text = get_text(para)
            if i < len(new_parts):
                replace(para, old_text, new_parts[i], tracked)
            else:
                clear_para(para)
        print(f"Summary: {len(sum_paras)} paras replaced", file=sys.stderr)

    # Replace Core Expertise — find table with bullet character
    expertise_items = [e.strip() for e in cv_data.get('expertise','').split(',') if e.strip()]
    half = len(expertise_items) // 2
    left_items = expertise_items[:half]
    right_items = expertise_items[half:]

    for table in doc.tables:
        if len(table.columns) == 2 and len(table.rows) == 1:
            c0, c1 = table.cell(0,0), table.cell(0,1)
            if '▸' in c0.text or '▸' in c1.text:
                lp = [p for p in c0.paragraphs if get_text(p)]
                rp = [p for p in c1.paragraphs if get_text(p)]
                for i, para in enumerate(lp):
                    new_t = '▸ ' + left_items[i] if i < len(left_items) else ''
                    runs = para.runs
                    if new_t:
                        if runs: runs[0].text = new_t; [r.__setattr__('text','') for r in runs[1:]]
                        else: para.add_run(new_t)
                    else:
                        clear_para(para)
                for i, para in enumerate(rp):
                    new_t = '▸ ' + right_items[i] if i < len(right_items) else ''
                    runs = para.runs
                    if new_t:
                        if runs: runs[0].text = new_t; [r.__setattr__('text','') for r in runs[1:]]
                        else: para.add_run(new_t)
                    else:
                        clear_para(para)
                print(f"Expertise: {len(left_items)} left, {len(right_items)} right", file=sys.stderr)
                break

    # Replace experience bullets
    if experience_idx is not None:
        all_new = []
        for exp in cv_data.get('experience', []):
            all_new.extend(exp.get('bullets', []))
        bullet_paras = [p for p in paras[experience_idx:] if p.style.name == 'List Paragraph' and get_text(p)]
        for i, para in enumerate(bullet_paras):
            if i >= len(all_new): break
            replace(para, get_text(para), all_new[i], tracked)
        print(f"Bullets: {min(len(bullet_paras), len(all_new))} replaced", file=sys.stderr)

    doc.save(output_path)
    print(f"Saved: {output_path}", file=sys.stderr)
    return True

if __name__ == '__main__':
    if len(sys.argv) < 4:
        print('Usage: cv_editor.py <input.docx> <output.docx> <cvdata.json> [tracked]', file=sys.stderr)
        sys.exit(1)
    input_path, output_path, cvdata_path = sys.argv[1], sys.argv[2], sys.argv[3]
    tracked = len(sys.argv) > 4 and sys.argv[4] == 'tracked'
    with open(cvdata_path) as f:
        cv_data = json.load(f)
    sys.exit(0 if edit_cv(input_path, output_path, cv_data, tracked) else 1)
